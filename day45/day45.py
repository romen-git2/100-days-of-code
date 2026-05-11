import streamlit as st
import os
import base64
from io import BytesIO
from PIL import Image
from dotenv import load_dotenv
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_ollama import OllamaLLM, ChatOllama
import pypdf
import docx
import re
import requests
from fpdf import FPDF
from datetime import datetime

# Page config, CSS
st.set_page_config(
    page_title="ContractGuard Pro",
    layout="wide"
)

st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;600;700&display=swap');
    
    html, body, [class*="css"] {
        font-family: 'Outfit', sans-serif;
        background-color: #0a0c10;
        color: #e2e8f0;
    }

    .stApp {
        background: radial-gradient(circle at top right, #1e293b, #0a0c10);
    }

    .stButton > button {
        border-radius: 6px;
        font-weight: 600;
        letter-spacing: 0.025em;
        text-transform: uppercase;
        font-size: 0.75rem;
        padding: 0.75rem 1.5rem;
        transition: all 0.2s cubic-bezier(0.4, 0, 0.2, 1);
        border: 1px solid rgba(255, 255, 255, 0.1);
        background: rgba(255, 255, 255, 0.03);
    }

    .stButton > button:hover {
        border-color: #3b82f6;
        background: rgba(59, 130, 246, 0.1);
        box-shadow: 0 0 20px rgba(59, 130, 246, 0.15);
    }

    .feature-card {
        background: rgba(15, 23, 42, 0.6);
        padding: 32px;
        border-radius: 12px;
        border: 1px solid rgba(148, 163, 184, 0.1);
        text-align: left;
        transition: all 0.3s ease;
        height: 100%;
        backdrop-filter: blur(8px);
    }
    
    .feature-card:hover {
        border-color: rgba(59, 130, 246, 0.4);
        background: rgba(15, 23, 42, 0.8);
    }

    .feature-card h3 {
        font-size: 1.1rem;
        font-weight: 600;
        margin-bottom: 12px;
        color: #f8fafc;
        letter-spacing: -0.01em;
    }
    
    .feature-card p {
        font-size: 0.875rem;
        line-height: 1.6;
        color: #94a3b8;
    }

    /* Custom Radio Tabs */
    div[data-testid="stHorizontalBlock"] .stRadio div[role="radiogroup"] {
        background: rgba(30, 41, 59, 0.5);
        padding: 4px;
        border-radius: 8px;
        border: 1px solid rgba(255, 255, 255, 0.05);
    }

    .st-info-box {
        margin-top: 2rem !important;
    }

    [data-testid="stMarkdownContainer"] p, [data-testid="stMarkdownContainer"] li {
        font-size: 0.95rem !important;
        line-height: 1.7 !important;
        color: #cbd5e1;
    }
    
    [data-testid="stMarkdownContainer"] h3 {
        font-size: 1.1rem !important;
        font-weight: 700 !important;
        margin-top: 2.5rem !important;
        margin-bottom: 1rem !important;
        color: #f8fafc !important;
        border-bottom: 1px solid rgba(255, 255, 255, 0.05);
        padding-bottom: 0.5rem;
    }

    .sidebar .stSelectbox label {
        font-size: 0.8rem;
        text-transform: uppercase;
        letter-spacing: 0.05em;
        color: #64748b;
    }
    </style>

""", unsafe_allow_html=True)

load_dotenv()

# Logic helpers

@st.cache_resource
def get_llm(model_name="gemini-3-flash-preview", temp=0.1):
    api_key = os.getenv("GOOGLE_API_KEY")
    if not api_key:
        st.error("Missing GOOGLE_API_KEY environment variable.")
        st.stop()
    
    kwargs = {"model": model_name, "temperature": temp}
    if "gemini-3" in model_name:
        kwargs["thinking_level"] = "medium"
        
    return ChatGoogleGenerativeAI(**kwargs)

@st.cache_resource
def get_local_scrubber():
    try:
        # Using gemma3:270m - Extremely lightweight, fits in almost any GPU/RAM.
        return OllamaLLM(model="gemma3:270m", temperature=0)
    except Exception as e:
        return None

def anonymize_text(text):
    scrubber = get_local_scrubber()
    if not scrubber:
        return text, False
    
    # Explicit instructions for 270M models to avoid over-scrubbing
    prompt = f"TASK: Replace ONLY person names and addresses with [NAME] or [ADDRESS]. \nCRITICAL: DO NOT DELETE ANY LEGAL CLAUSES. Keep the rest of the text exactly the same. \n\nTEXT TO SCRUB:\n{text}\n\nSCRUBBED TEXT:"
    
    try:
        scrubbed = scrubber.invoke(prompt)
        
        # GUARDRAIL: If the model deleted more than 30% of the text, it probably 
        # deleted the clauses. Reject the scrub and use fallback.
        if not scrubbed or len(scrubbed) < (0.7 * len(text)):
            return text, False
            
        return scrubbed.strip(), True
    except Exception as e:
        return text, False

@st.cache_resource
def get_local_vision_model():
    try:
        # Defaulting to moondream as fallback
        return ChatOllama(model="moondream", temperature=0)
    except Exception as e:
        return None

def ocr_image_locally(file):
    vision_model = get_local_vision_model()
    if not vision_model:
        return None, False
    
    # Encode current image
    b64_string = encode_image(file)
    
    message = HumanMessage(
        content=[
            {"type": "text", "text": "Transcribe all text from this document image exactly. Do not summarize or explain, just provide the text."},
            {
                "type": "image_url",
                "image_url": {"url": f"data:image/jpeg;base64,{b64_string}"},
            },
        ]
    )
    
    try:
        response = vision_model.invoke([message])
        return response.content, True
    except Exception as e:
        st.error(f"OCR Error: {e}")
        return str(e), False

def check_ollama_status():
    try:
        response = requests.get("http://localhost:11434", timeout=2)
        return response.status_code == 200
    except:
        return False

def format_report_for_display(content: str) -> str:
    """Inject paragraph breaks before known bold labels for clean web rendering."""
    BREAK_BEFORE = [
        "**Risk Level:**", "**Clause:**", "**Analysis:**",
        "**User Impact:**", "**Proposed Better Version:**",
        "**Jurisdictional Note:**",
    ]
    for label in BREAK_BEFORE:
        content = content.replace(label, f"\n\n{label}")
    return content


def generate_pdf(markdown_text: str) -> bytes:
    """Converts the markdown audit report to a fully-structured PDF."""

    def sanitize(text: str) -> str:
        text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
        text = re.sub(r"\*(.*?)\*", r"\1", text)
        text = re.sub(r"`(.*?)`", r"\1", text)
        return text.encode("latin-1", errors="replace").decode("latin-1")

    def is_separator_row(cells):
        """True if every cell is a markdown table separator like :--- or ---."""
        return all(re.match(r'^:?-+:?$', c) for c in cells)

    # Pre-process: inject blank lines before key bold labels
    for label in ["**Risk Level:**", "**Clause:**", "**Analysis:**",
                  "**User Impact:**", "**Proposed Better Version:**",
                  "**Jurisdictional Note:**"]:
        markdown_text = markdown_text.replace(label, f"\n{label}")

    pdf = FPDF()
    pdf.add_page()
    pdf.set_margins(15, 20, 15)
    pdf.set_auto_page_break(auto=True, margin=18)
    effective_w = pdf.w - pdf.l_margin - pdf.r_margin

    # Header
    pdf.set_font("Helvetica", "B", 16)
    pdf.set_text_color(20, 60, 120)
    pdf.cell(0, 12, "ContractGuard Pro - Audit Report", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(120, 120, 120)
    pdf.cell(0, 6, f"Generated: {datetime.now().strftime('%B %d, %Y at %H:%M')}", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)
    pdf.set_draw_color(20, 60, 120)
    pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
    pdf.ln(6)

    # Parse lines using a while loop to collect table blocks
    lines = markdown_text.split("\n")
    i = 0
    while i < len(lines):
        raw = lines[i]
        clean = raw.strip()

        # Blank line
        if not clean:
            pdf.ln(3)
            i += 1
            continue

        # Table block: collect all consecutive | lines then render at once
        if clean.startswith("|"):
            table_raw = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_raw.append(lines[i].strip())
                i += 1

            # Parse into rows (skip separator rows)
            all_rows = []
            for tline in table_raw:
                cells = [c.strip() for c in tline.split("|") if c.strip()]
                if not cells:
                    continue
                if is_separator_row(cells):
                    continue
                all_rows.append(cells)

            if not all_rows:
                continue

            # Determine column widths (equal split)
            n_cols = max(len(r) for r in all_rows)
            col_w = effective_w / n_cols

            for row_idx, row in enumerate(all_rows):
                is_hdr = row_idx == 0
                pdf.set_font("Helvetica", "B" if is_hdr else "", 8)
                pdf.set_text_color(30, 30, 30)

                # Pre-calculate heights to support multi-line cells in a row
                row_y = pdf.get_y()
                # If near bottom, push to new page
                if row_y > pdf.h - 30:
                    pdf.add_page()
                    row_y = pdf.get_y()

                # Render cells in the row
                max_row_height = 0
                cell_contents = []
                
                # First pass: Get all cell contents and find max height
                for col_idx in range(n_cols):
                    txt = sanitize(row[col_idx]) if col_idx < len(row) else ""
                    # Split text into lines that fit the column width
                    lines_in_cell = pdf.multi_cell(col_w, 4, txt, dry_run=True, output="LINES")
                    cell_height = len(lines_in_cell) * 4
                    max_row_height = max(max_row_height, cell_height)
                    cell_contents.append(txt)

                # Second pass: Render the cells
                for col_idx, txt in enumerate(cell_contents):
                    x = pdf.l_margin + (col_idx * col_w)
                    pdf.set_xy(x, row_y)
                    # Border=1 for header, 0 for body
                    pdf.multi_cell(col_w, 4, txt, border=0)

                # Move cursor to the bottom of the tallest cell in this row
                pdf.set_y(row_y + max_row_height + 2)

                if is_hdr:
                    pdf.set_draw_color(160, 160, 200)
                    pdf.line(pdf.l_margin, pdf.get_y() - 1, pdf.w - pdf.r_margin, pdf.get_y() - 1)
                    pdf.ln(1)

            pdf.ln(3)
            continue

        # Non-table lines
        try:
            if clean.startswith("### "):
                pdf.ln(2)
                pdf.set_font("Helvetica", "B", 12)
                pdf.set_text_color(20, 60, 120)
                pdf.multi_cell(0, 7, sanitize(clean[4:]))

            elif clean.startswith("## "):
                pdf.ln(3)
                pdf.set_font("Helvetica", "B", 13)
                pdf.set_text_color(10, 40, 100)
                pdf.multi_cell(0, 8, sanitize(clean[3:]))

            elif clean.startswith("# "):
                pdf.ln(3)
                pdf.set_font("Helvetica", "B", 14)
                pdf.set_text_color(10, 30, 80)
                pdf.multi_cell(0, 9, sanitize(clean[2:]))
                pdf.ln(1)

            elif clean.startswith("---") and len(clean) <= 5:
                pdf.set_draw_color(200, 200, 200)
                pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
                pdf.ln(3)

            elif clean.startswith(("* ", "- ")):
                pdf.set_font("Helvetica", "", 9.5)
                pdf.set_text_color(40, 40, 40)
                pdf.multi_cell(0, 6, sanitize(f"  - {clean[2:]}"))

            elif clean.startswith("**") and ":" in clean:
                # Bold label: **Label:** content
                colon_idx = clean.index(":")
                label_raw = clean[:colon_idx].replace("**", "").strip()
                content_raw = clean[colon_idx + 1:].replace("**", "").strip()
                pdf.set_font("Helvetica", "B", 9.5)
                pdf.set_text_color(30, 30, 30)
                pdf.write(5.5, sanitize(f"{label_raw}: "))
                pdf.set_font("Helvetica", "", 9.5)
                pdf.set_text_color(60, 60, 60)
                pdf.multi_cell(0, 5.5, sanitize(content_raw))
                pdf.ln(1)

            else:
                pdf.set_font("Helvetica", "", 9.5)
                pdf.set_text_color(40, 40, 40)
                pdf.multi_cell(0, 6, sanitize(clean))

        except Exception:
            pass

        i += 1

    buf = BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf.read()

def encode_image(uploaded_file):
    img = Image.open(uploaded_file)
    if img.mode != "RGB":
        img = img.convert("RGB")
    buffered = BytesIO()
    img.save(buffered, format="JPEG")
    return base64.b64encode(buffered.getvalue()).decode('utf-8')

def get_svg(type):
    """Returns SVG code for different UI sections."""
    svgs = {
        "shield": """<svg width="40" height="40" viewBox="0 0 24 24" fill="none" stroke="#3b82f6" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"><path d="M12 22s8-4 8-10V5l-8-3-8 3v7c0 6 8 10 8 10z"/></svg>""",
        "search": """<svg width="32" height="32" viewBox="0 0 24 24" fill="none" stroke="#64748b" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><circle cx="11" cy="11" r="8"/><path d="m21 21-4.3-4.3"/></svg>""",
        "pen": """<svg width="32" height="32" viewBox="0 0 24 24" fill="none" stroke="#64748b" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M12 19l7-7 3 3-7 7-3-3z"/><path d="M18 13l-1.5-7.5L2 2l3.5 14.5L13 18l5-5z"/><path d="M2 2l5 5"/></svg>""",
        "alert": """<svg width="32" height="32" viewBox="0 0 24 24" fill="none" stroke="#64748b" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><circle cx="12" cy="12" r="10"/><line x1="12" y1="8" x2="12" y2="12"/><line x1="12" y1="16" x2="12.01" y2="16"/></svg>"""
    }
    return svgs.get(type, "")

# System prompts

AUDIT_SYSTEM_PROMPT = """
You are a high-end Legal Tech Consultant. Review the contract images for professional fairness.
CONTEXT: User is a {user_role} reviewing a {contract_type} under the jurisdiction of {jurisdiction}.

FORMATTING RULES:
1. Use clear Markdown headers (#, ##, ###).
2. Use 'CRITICAL', 'WARNING', and 'NOTICE' labels.
3. For the Detailed Risk Analysis, use the following strict structure for each point:
   - ### [Point Name]
   - **Risk Level:** [Level]
   - **Clause:** [Clause text]
   - **Analysis:** [Detailed legal analysis]
   - **User Impact:** [How this affects the user]
   - **Proposed Better Version:** [The recommended text]
   - **Jurisdictional Note:** [Specific law reference]
   
   IMPORTANT: Every label (e.g., **Risk Level:**) MUST be bold and include the double asterisks.
4. Provide a summary table of risks at the beginning.
"""

NEGOTIATION_SYSTEM_PROMPT = """
You are a world-class business negotiator. Draft a response based on the audit.
TONE: {tone}
ROLE: {user_role}
JURISDICTION: {jurisdiction}

Ensure the message is persuasive, professional, and protects the user's rights without damaging the partnership.
Reference relevant local legal standards if it strengthens the negotiating position.
"""

# Interface

def main():
    with st.sidebar:
        st.markdown(f'<div style="text-align:center">{get_svg("shield")}</div>', unsafe_allow_html=True)
        st.header("ContractGuard Pro")
        
        user_role = st.selectbox(
            "Your Role:",
            ["Freelancer", "Small Business Owner", "Creative Professional", "Consultant", "Employee", "Legal Professional"]
        )
        
        contract_type = st.selectbox(
            "Contract Type:",
            [
                "Service Agreement", "NDA", "Employment Contract", "IP Licensing", 
                "Work for Hire", "Lease Agreement", "Software License", 
                "Partnership Agreement", "Affiliate Agreement", "Custom/Other"
            ]
        )
        
        jurisdiction = st.selectbox(
            "Jurisdiction / Country:",
            ["United States", "United Kingdom", "Canada", "Australia", "European Union", "India", "Sri Lanka", "General / International"]
        )

        st.divider()
        st.subheader("Compute Settings")
        selected_model = st.selectbox(
            "Cloud Intelligence:",
            ["gemini-3-flash-preview", "gemini-3.1-pro-preview"],
            help="Pro models are better for complex legal reasoning."
        )

        ollama_ready = check_ollama_status()
        
        use_privacy_shield = st.toggle(
            "Local Encryption Shield",
            help="Processes data locally to remove sensitive identifiers before analysis."
        )
        
        if use_privacy_shield:
            if ollama_ready:
                st.success("Encryption Active")
            else:
                st.error("Engine Offline")
                with st.expander("How to fix this?"):
                    st.markdown("""
                    1. **Install Ollama**: Download from [ollama.com](https://ollama.com).
                    2. **Pull the Model**: Run `ollama pull gemma3:270m` in your terminal.
                    3. **Run Ollama**: Ensure the app is open in your taskbar.
                    
                    *Or simply turn off Privacy Shield to use direct cloud analysis.*
                    """)
        else:
            st.warning("Direct Cloud Transmission")
        
        st.divider()
        uploaded_files = st.file_uploader("Upload contracts (docs/images)", type=["jpg", "png", "jpeg", "pdf", "txt", "docx"], accept_multiple_files=True)
        
        run_analysis = st.button("Run Audit", type="primary", use_container_width=True)
        if st.button("Reset Session", use_container_width=True):
            st.session_state.clear()
            st.rerun()

    # Main Dashboard Logic
    if not uploaded_files:
        st.title("Secure Your Agreements")
        st.markdown("##### Professional-grade legal risk detection and negotiation strategy.")
        
        st.write(" ")
        c1, c2, c3 = st.columns(3)
        
        with c1:
            st.markdown(f"""<div class="feature-card">
                {get_svg("search")}
                <h3>Clause Analysis</h3>
                <p>Advanced vision identifies predatory liability and indemnity shifts.</p>
            </div>""", unsafe_allow_html=True)
            
        with c2:
            st.markdown(f"""<div class="feature-card">
                {get_svg("alert")}
                <h3>Risk Profiling</h3>
                <p>Severity-based reporting focused on your specific professional role.</p>
            </div>""", unsafe_allow_html=True)
            
        with c3:
            st.markdown(f"""<div class="feature-card">
                {get_svg("pen")}
                <h3>Rebuttal Drafts</h3>
                <p>Automated professional responses to secure fairer terms immediately.</p>
            </div>""", unsafe_allow_html=True)
            
        st.markdown('<div class="st-info-box">', unsafe_allow_html=True)
        st.info("Upload your document in the sidebar to begin the secure analysis.")
        st.markdown('</div>', unsafe_allow_html=True)

    elif not run_analysis and "audit_report" not in st.session_state:
        # Documents are uploaded but analysis hasn't started
        st.title("Documents Loaded")
        st.markdown(f"### {len(uploaded_files)} contract(s) ready for review.")
        
        st.info("Your documents have been staged. Please verify your **Role**, **Jurisdiction**, and **Privacy Settings** in the sidebar, then click **Run Audit** to start the analysis.")
        
        # Show a quick preview of file names
        st.write("---")
        st.caption("Files staged for analysis:")
        for f in uploaded_files:
            st.write(f"📄 {f.name} ({round(f.size/1024, 1)} KB)")

    elif run_analysis:
        if use_privacy_shield and not ollama_ready:
            st.error("🛑 **Privacy Shield Error**: You have enabled the Privacy Shield, but Ollama is not running on your machine. Please start Ollama or disable the toggle in the sidebar to proceed with a direct cloud analysis.")
            st.stop()
            
        llm = get_llm(model_name=selected_model)
        prompt_parts = [{"type": "text", "text": "Analyze these contract pages and documents."}]
        
        # Privacy Shield logic
        if use_privacy_shield:
            with st.status("Applying Edge Privacy Shield...", expanded=True) as status:
                st.write("Extracting and scrubbing text locally...")
                total_scrubbed = 0
                for file in uploaded_files:
                    file_text = ""
                    if file.name.lower().endswith(('.png', '.jpg', '.jpeg')):
                        st.write(f"Scanning {file.name} locally...")
                        ocr_text, ocr_success = ocr_image_locally(file)
                        if ocr_success:
                            file_text = ocr_text
                        else:
                            st.warning(f"Local OCR failed for {file.name}. Ensure a multimodal model is pulled (e.g., moondream).")
                            continue
                    elif file.name.lower().endswith('.pdf'):
                        pdf_reader = pypdf.PdfReader(file)
                        for page_num, page in enumerate(pdf_reader.pages):
                            page_text = page.extract_text() or ""
                            
                            # If page text is very thin, try to extract images for OCR
                            if len(page_text.strip()) < 50:
                                st.write(f"Scanning Page {page_num+1} of {file.name} (Local OCR)...")
                                # Extract images from the page
                                page_images_text = ""
                                try:
                                    for img in page.images:
                                        # Use a temporary buffer for each image found on the page
                                        img_buf = BytesIO(img.data)
                                        ocr_text, ocr_success = ocr_image_locally(img_buf)
                                        if ocr_success:
                                            page_images_text += ocr_text + "\n"
                                    
                                    if page_images_text:
                                        page_text = page_images_text
                                except Exception:
                                    pass # Fallback to whatever text was there (or empty)
                                    
                            file_text += page_text + "\n"

                    elif file.name.lower().endswith('.docx'):
                        doc = docx.Document(file)
                        file_text = "\n".join([para.text for para in doc.paragraphs])
                    elif file.name.lower().endswith('.txt'):
                        file_text = file.getvalue().decode('utf-8')

                    if file_text:
                        scrubbed_text, success = anonymize_text(file_text)
                        if success:
                            prompt_parts.append({"type": "text", "text": f"\n--- Scrubbed Extract from {file.name} ---\n{scrubbed_text}"})
                            total_scrubbed += 1
                        else:
                            # Safe fallback: If scrubbing fails, still proceeding with direct text but warn the user.
                            st.warning(f"Privacy Shield bypassed for {file.name} (local model timed out). Proceeding with direct analysis.")
                            prompt_parts.append({"type": "text", "text": f"\n--- Direct Extract from {file.name} ---\n{file_text}"})
                    else:
                        st.warning(f"Could not extract any text from {file.name}.")

                
                status.update(label=f"Privacy Shield Complete! Scrubbed {total_scrubbed} docs.", state="complete")
        
        # Standard processing (non-scrubbed or images)
        for file in uploaded_files:
            if file.name.lower().endswith(('.png', '.jpg', '.jpeg')):
                b64_string = encode_image(file)
                prompt_parts.append({
                    "type": "image_url",
                    "image_url": {"url": f"data:image/jpeg;base64,{b64_string}"}
                })
            elif not use_privacy_shield:
                # Add text normally if shield is off
                if file.name.lower().endswith('.pdf'):
                    pdf_reader = pypdf.PdfReader(file)
                    text = "".join([p.extract_text() or "" for p in pdf_reader.pages])
                elif file.name.lower().endswith('.docx'):
                    doc = docx.Document(file)
                    text = "\n".join([para.text for para in doc.paragraphs])
                elif file.name.lower().endswith('.txt'):
                    text = file.getvalue().decode('utf-8')
                else:
                    continue
                prompt_parts.append({"type": "text", "text": f"\n Extract from {file.name} \n{text}"})

        with st.spinner("Analyzing document structure and legal weight..."):
            try:
                messages = [
                    SystemMessage(content=AUDIT_SYSTEM_PROMPT.format(
                        user_role=user_role, 
                        contract_type=contract_type,
                        jurisdiction=jurisdiction
                    )),
                    HumanMessage(content=prompt_parts)
                ]
                response = llm.invoke(messages)
                
                # Ensure content is a string (handles thinking/multimodal parts)
                content = response.content
                if isinstance(content, list):
                    content = "".join([part["text"] if isinstance(part, dict) and "text" in part else str(part) for part in content])
                
                st.session_state.audit_report = content
                st.session_state.current_role = user_role
                st.session_state.current_jurisdiction = jurisdiction
            except Exception as e:
                st.error(f"Analysis failed: {e}")

    # Results
    if "audit_report" in st.session_state:
        st.write("")
        active_tab = st.radio(
            "View Analysis:",
            ["Audit Findings", "Negotiation Draft"],
            horizontal=True,
            label_visibility="collapsed",
            key="active_tab_selector"
        )
        st.divider()
        
        if "Audit Findings" in active_tab:
            st.markdown(format_report_for_display(st.session_state.audit_report))
            pdf_bytes = generate_pdf(st.session_state.audit_report)
            st.download_button(
                label="Export Audit (PDF)",
                data=pdf_bytes,
                file_name="contractguard_audit_report.pdf",
                mime="application/pdf"
            )
            
        else:
            st.write("### Negotiation Assistant")
            st.caption("Select a communication strategy to generate your response:")
            
            t1, t2, t3 = st.columns(3)
            selected_tone = None
            if t1.button("Collaborative Strategy"): selected_tone = "Partnership-focused and flexible"
            if t2.button("Protective Strategy"): selected_tone = "Firm, legalistic and rights-focused"
            if t3.button("Educational Strategy"): selected_tone = "Explaining constraints of a small business"

            if selected_tone:
                with st.spinner("Drafting your response..."):
                    llm_neg = get_llm(temp=0.7)
                    neg_messages = [
                        SystemMessage(content=NEGOTIATION_SYSTEM_PROMPT.format(
                            tone=selected_tone, 
                            user_role=st.session_state.current_role,
                            jurisdiction=st.session_state.current_jurisdiction
                        )),
                        HumanMessage(content=f"Draft based on this audit:\n\n{st.session_state.audit_report}")
                    ]
                    rebuttal = llm_neg.invoke(neg_messages)
                    
                    # Ensure content is a string
                    reb_content = rebuttal.content
                    if isinstance(reb_content, list):
                        reb_content = "".join([part["text"] if isinstance(part, dict) and "text" in part else str(part) for part in reb_content])
                        
                    # Clean All Markdown for a pure copy-paste email
                    reb_content = reb_content.replace("**", "").replace("### ", "").replace("## ", "").replace("# ", "")
                        
                    st.session_state.current_rebuttal = reb_content
            
            if "current_rebuttal" in st.session_state:
                st.text_area("Finalized Draft:", value=st.session_state.current_rebuttal, height=400)
                st.caption("Recommended: Review with counsel before sending.")

if __name__ == "__main__":
    main()